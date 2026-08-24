import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("🛡️ AegisSOC Global Distributed Multi-Node SIEM System\nComplete End-to-End Project Documentation")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x06, 0x4E, 0x8A)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Verified Live Multi-Node Demonstration Gallery | WIN-RJCSG0BSFFA (192.168.0.181) -> SIEM Server (192.168.0.128:8000)")
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph() # Spacer

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Verified Live Visual Gallery", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x06, 0x4E, 0x8A)
    
    doc.add_paragraph(
        "AegisSOC is a multi-node distributed Security Information and Event Management (SIEM) and Security Operations Center (SOC) "
        "dashboard built specifically to collect, correlate, and monitor Windows host logs and telemetry across multiple remote PCs on a network."
    )

    doc.add_heading("Verified Live Network Screenshots (5-Image Gallery):", level=2)
    doc.add_paragraph(
        "1. Remote Agent Terminal (agent.py): Live execution on remote target WIN-RJCSG0BSFFA (192.168.0.181) streaming Event IDs 5156, 1102, 4624, 4688.\n"
        "2. Command Center HUD: 4 Fleet Nodes Active, 203+ EPS ingestion velocity, 88% Fleet Health Score, and Lateral Movement Warning Banner.\n"
        "3. SIEM Log Workbench: Live streaming log grid displaying 38+ real-time events across WIN-RJCSG0BSFFA and fleet nodes.\n"
        "4. Registered Fleet Management: Active fleet cards displaying WIN-RJCSG0BSFFA (Friend PC) registered live alongside fleet nodes.\n"
        "5. MITRE ATT&CK Matrix & Incidents: Automated threat correlation for WIN-RJCSG0BSFFA under TA0002 Execution (PowerShell Execution T1059.001) & TA0005 Defense Evasion (Clear Event Logs T1070.001) with 1-click Host Isolation playbooks."
    )

    output_path = os.path.abspath("AegisSOC_Complete_Project_Guide.docx")
    doc.save(output_path)
    print(f"Updated Word document saved at: {output_path}")

if __name__ == "__main__":
    create_document()
